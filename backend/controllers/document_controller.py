import os
import re
import fitz
import docx
import uuid
import pytesseract
import requests
from pdf2image import convert_from_path
from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException
from sqlalchemy.orm import Session
from pydantic import BaseModel
from datetime import datetime
from openai import OpenAI
from PIL import Image

from utils.auth import get_current_user
from models.user import User
from models.notebook import Notebook
from models.document import Document
from database import get_db
from utils.chunking import chunk_text
from utils.embedding import get_embedding
from utils.vectorstore import collection, add_chunk
from utils.youtube_helper import get_youtube_transcript, extract_video_id
from langchain_text_splitters import RecursiveCharacterTextSplitter

router = APIRouter()

# Thêm biến UPLOAD_ROOT để tránh lỗi undefined ở API Upload
UPLOAD_ROOT = "./storage"
os.makedirs(UPLOAD_ROOT, exist_ok=True)

class YouTubeRequest(BaseModel):
    url: str
    notebook_id: str = "default"

class TextUploadRequest(BaseModel):
    notebook_id: int  
    text: str

def extract_video_id(url: str) -> str:
    """Hàm tự chế để bóc tách ID video từ link YouTube"""
    match = re.search(r"(?:v=|\/)([0-9A-Za-z_-]{11}).*", url)
    if match:
        return match.group(1)
    return "unknown_video_id"

@router.get("/notebooks/{notebook_id}/documents/")
async def get_documents(
    notebook_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    nb = db.query(Notebook).filter(Notebook.id == notebook_id, Notebook.user_id == current_user.id).first()
    if not nb:
        raise HTTPException(status_code=404, detail="Notebook không tồn tại hoặc không thuộc quyền sở hữu.")

    docs = db.query(Document).filter(Document.notebook_id == notebook_id).all()
    
    return {
        "documents": [
            {
                "id": doc.id, 
                "name": doc.filename, 
                "type": doc.filetype
            } 
            for doc in docs
        ]
    }

# ==========================================
# API UPLOAD FILE PDF / DOCX / TXT (Chính)
# ==========================================
@router.post("/notebooks/{notebook_id}/upload/")
async def upload_document(
    notebook_id: int,
    file: UploadFile = File(...), 
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 1. CHUẨN HÓA TÊN FILE NGAY LẬP TỨC
    clean_filename = file.filename.lower().strip()
    
    ext = clean_filename.split('.')[-1]
    if ext not in ["pdf", "docx", "txt", "png", "jpg", "jpeg"]:
        return {"error": "Chỉ hỗ trợ .pdf, .docx, .txt và hình ảnh (.png, .jpg, .jpeg)"}

    nb = db.query(Notebook).filter(Notebook.id == notebook_id, Notebook.user_id == current_user.id).first()
    if not nb:
        return {"error": "Notebook không hợp lệ."}

    current_docs = db.query(Document).filter(Document.notebook_id == notebook_id).count()
    if current_docs >= 10:
        return {"error": "Notebook đã đạt giới hạn tối đa 10 tài liệu."}

    # 2. KIỂM TRA TRÙNG LẶP THEO TÊN ĐÃ CHUẨN HÓA
    existing_doc = db.query(Document).filter(
        Document.notebook_id == notebook_id, 
        Document.filename == clean_filename
    ).first()
    
    if existing_doc:
        return {"error": "Tài liệu này đã tồn tại trong Notebook này. Vui lòng đổi tên file khác."}

    # 3. LƯU FILE TẠM VÀ TRÍCH XUẤT VĂN BẢN
    temp_path = os.path.join(UPLOAD_ROOT, file.filename)
    with open(temp_path, "wb") as f:
        f.write(await file.read())

    text = ""
    if ext == "pdf":
        doc = fitz.open(temp_path)
        text = "".join([page.get_text() for page in doc])
        if len(text.strip()) < 50:
            print(f"Phát hiện PDF Scan ({clean_filename}). Đang kích hoạt Tesseract OCR...")
            try:
                images = convert_from_path(temp_path)
                ocr_text = []
                for img in images:
                    ocr_text.append(pytesseract.image_to_string(img, lang='vie'))
                text = "\n".join(ocr_text)
            except Exception as e:
                os.remove(temp_path)
                print(f"Lỗi OCR: {str(e)}")
                return {"error": "Hệ thống thiếu thư viện xử lý ảnh (Poppler/Tesseract) để đọc file Scan."}
    elif ext == "docx":
        doc = docx.Document(temp_path)
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip(): text_parts.append(p.text)
        for table in doc.tables:
            text_parts.append("\n[BẮT ĐẦU BẢNG]")
            for row in table.rows:
                row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                text_parts.append(" | ".join(row_data))
            text_parts.append("[KẾT THÚC BẢNG]\n")
        text = "\n".join(text_parts)
        
    # ==================================================
    # LOGIC OCR XỬ LÝ ẢNH ĐƯỢC CHÈN VÀO ĐÂY
    # ==================================================
    elif ext in ["png", "jpg", "jpeg"]:
        print(f"Đang kích hoạt Tesseract OCR cho ảnh: {clean_filename}...")
        try:
            img = Image.open(temp_path)
            # Dùng lang='vie+eng' để đọc mượt cả Tiếng Việt và Tiếng Anh
            text = pytesseract.image_to_string(img, lang='vie+eng') 
        except Exception as e:
            os.remove(temp_path)
            print(f"Lỗi OCR Ảnh: {str(e)}")
            return {"error": "Không thể đọc được chữ từ hình ảnh này."}
    # ==================================================
    
    else:
        # Nếu là file .txt thì vào đây
        with open(temp_path, encoding="utf-8") as f:
            text = f.read()

    if not text.strip():
        os.remove(temp_path)
        return {"error": "Không thể trích xuất được văn bản nào từ file này."}

    # 4. LƯU SQL BẰNG TÊN ĐÃ CHUẨN HÓA
    new_doc = Document(
        notebook_id=notebook_id, 
        filename=clean_filename, 
        filetype=ext,
        content=text
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    # 5. LƯU CHROMADB BẰNG TÊN ĐÃ CHUẨN HÓA
    chunks = chunk_text(text, chunk_size=1000, overlap=200) 
    for i, chunk in enumerate(chunks):
        chunk_id = f"{notebook_id}_{clean_filename}_{i}" 
        embedding = get_embedding(chunk)
        add_chunk(chunk_id, chunk, embedding, metadata={"notebook_id": int(notebook_id), "filename": clean_filename})

    os.remove(temp_path)

    return {"message": f"Upload thành công file {clean_filename} ({len(chunks)} chunks)"}

# ==========================================
# XEM TRƯỚC VÀ XÓA TÀI LIỆU
# ==========================================
@router.get("/documents/{identifier}")
async def get_document_preview(identifier: str, db: Session = Depends(get_db)):
    """API xem trước thông minh: Chấp nhận cả ID (số) hoặc Tên file (chuỗi)"""
    if identifier.isdigit():
        doc = db.query(Document).filter(Document.id == int(identifier)).first()
    else:
        # CHUẨN HÓA TEXT TỪ GIAO DIỆN GỬI LÊN
        clean_identifier = identifier.lower().strip()
        doc = db.query(Document).filter(Document.filename == clean_identifier).first()
        
    if not doc:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu")
        
    return {
        "id": doc.id,
        "filename": doc.filename,
        "filetype": doc.filetype,
        "content": doc.content or "Tài liệu này không có nội dung văn bản thô để hiển thị."
    }

@router.delete("/notebooks/{notebook_id}/documents/{identifier}")
async def delete_document(
    notebook_id: int,
    identifier: str, 
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    nb = db.query(Notebook).filter(Notebook.id == notebook_id, Notebook.user_id == current_user.id).first()
    if not nb:
        raise HTTPException(status_code=404, detail="Notebook không hợp lệ.")

    if identifier.isdigit():
        doc = db.query(Document).filter(
            Document.id == int(identifier), 
            Document.notebook_id == notebook_id
        ).first()
    else:
        # CHUẨN HÓA TEXT TỪ GIAO DIỆN GỬI LÊN
        clean_identifier = identifier.lower().strip()
        doc = db.query(Document).filter(
            Document.filename == clean_identifier, 
            Document.notebook_id == notebook_id
        ).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu này.")

    filename_to_delete = doc.filename

    try:
        file_path = os.path.join(UPLOAD_ROOT, filename_to_delete)
        if os.path.exists(file_path):
            os.remove(file_path)

        collection.delete(
            where={
                "$and": [
                    {"notebook_id": notebook_id},
                    {"filename": filename_to_delete}
                ]
            }
        )
        
        db.delete(doc)
        db.commit()
        
        return {"message": f"Đã xóa thành công tài liệu: {filename_to_delete}"}
    
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Lỗi khi xóa tài liệu: {str(e)}")

# ==========================================
# YOUTUBE VÀ TEXT
# ==========================================
def get_youtube_title(url: str) -> str:
    try:
        response = requests.get(f"https://www.youtube.com/oembed?url={url}&format=json", timeout=5)
        if response.status_code == 200:
            return response.json().get("title", "YouTube Video")
    except Exception as e:
        print(f"Lỗi khi lấy tên YouTube: {e}")
    return "YouTube Video" 

@router.post("/upload-youtube")
async def upload_youtube(
    request: YouTubeRequest, 
    db: Session = Depends(get_db) 
):
    try:
        transcript_text = get_youtube_transcript(request.url)
        video_title = get_youtube_title(request.url)
        
        # CHUẨN HÓA TÊN FILE YOUTUBE
        clean_filename = f"[youtube] {video_title}".lower().strip()
        filetype = "youtube"
        
        new_doc = Document(
            notebook_id=request.notebook_id,
            filename=clean_filename, 
            filetype=filetype,
            content=transcript_text
        )
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)
        
        chunks = chunk_text(transcript_text, chunk_size=1000, overlap=200)
        chunk_ids = []
        for i, chunk in enumerate(chunks):
            chunk_id = f"{request.notebook_id}_{clean_filename}_{i}"
            embedding = get_embedding(chunk)
            
            metadata = {
                "chunk_index": i,
                "notebook_id": int(request.notebook_id),
                "filename": clean_filename,
                "source_type": filetype
            }
            add_chunk(chunk_id, chunk, embedding, metadata)
            chunk_ids.append(chunk_id)

        return {
            "status": "success", 
            "message": f"Đã học xong video! Tạo ra {len(chunks)} mảnh kiến thức.",
            "preview": transcript_text[:100] + "..."
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@router.post("/upload-text")
async def upload_pasted_text(request: TextUploadRequest, db: Session = Depends(get_db)):
    try:
        if not request.text.strip():
            return {"error": "Văn bản không được để trống!"}

        time_suffix = datetime.now().strftime("%H%M%S")
        
        # CHUẨN HÓA TÊN FILE TEXT
        clean_filename = f"pasted text ({time_suffix})".lower().strip()
        filetype = "text"
        
        new_doc = Document(
            notebook_id=request.notebook_id,
            filename=clean_filename,
            filetype=filetype,
            content=request.text 
        )
        db.add(new_doc)
        db.commit()
        db.refresh(new_doc)

        chunks = chunk_text(request.text, chunk_size=1000, overlap=200)
        if not chunks or len(chunks) == 0:
            chunks = [request.text]

        chunk_ids = []
        for i, chunk in enumerate(chunks):
            chunk_id = f"{request.notebook_id}_{clean_filename}_{i}"
            embedding = get_embedding(chunk)
            
            metadata = {
                "chunk_index": i,
                "notebook_id": int(request.notebook_id),
                "filename": clean_filename,
                "source_type": filetype
            }
            
            add_chunk(chunk_id, chunk, embedding, metadata)
            chunk_ids.append(chunk_id)

        return {
            "status": "success", 
            "message": "Đã lưu đoạn văn bản và Vector thành công!",
            "document_id": new_doc.id,
            "filename": clean_filename,
            "chunks": len(chunks)
        }
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(e))